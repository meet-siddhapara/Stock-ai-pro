import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set page margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69) # A4

# Add page numbers to footer
def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

for section in doc.sections:
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_para.add_run())

    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "[Ahmedabad University Logo]  |  Stock price prediction system"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Define Styles
styles = doc.styles

normal = styles['Normal']
normal.font.name = 'Arial'
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.5

h1 = styles['Heading 1']
h1.font.name = 'Arial'
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 51, 153)

h2 = styles['Heading 2']
h2.font.name = 'Arial'
h2.font.size = Pt(14)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 51, 153)

h3 = styles['Heading 3']
h3.font.name = 'Arial'
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0, 51, 153)

# 1. Cover Page
doc.add_paragraph('\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
runner = p.add_run('Stock price prediction system\n')
runner.font.size = Pt(24)
runner.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
runner2 = p2.add_run('A Project Report\n\n')
runner2.font.size = Pt(16)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
runner3 = p3.add_run('Submitted in partial fulfillment of the requirements for the degree of\n')
runner3.font.size = Pt(12)
p3.add_run('Bachelor of Technology in CSE\n\n').bold = True

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.add_run('By\n')
p4.add_run('meet siddhapara (AU2240243)\n\n').bold = True

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
p5.add_run('\n[Ahmedabad University Logo]\n\n').bold = True

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
run6 = p6.add_run('School of Engineering and Applied Science (SEAS)\nAhmedabad University\nApril, 2026')
run6.font.size = Pt(14)

doc.add_page_break()

# 2. Declaration Page
doc.add_heading('Declaration', level=1)
p = doc.add_paragraph("I hereby declare that the project report entitled \"Stock price prediction system\" submitted by me for the partial fulfillment of the requirements for the degree of Bachelor of Technology in CSE to the School of Engineering and Applied Science (SEAS), Ahmedabad University, is an authentic record of my own original work.")
p.paragraph_format.line_spacing = 1.5
doc.add_paragraph("I further declare that this work has not been submitted in full or in part to any other institution or university for the award of any degree or diploma.")
doc.add_paragraph('\n\n\n')
p_sig = doc.add_paragraph("Date: ____________________\t\t\t\tSignature: ____________________")
p_sig2 = doc.add_paragraph("Place: ____________________\t\t\t\tName: meet siddhapara")
doc.add_page_break()

# 3. Industry Mentor Certificate
doc.add_heading('Industry Mentor Certificate', level=1)
doc.add_paragraph("This is to certify that the project entitled \"Stock price prediction system\" has been carried out by meet siddhapara (AU2240243) as part of their internship at Grownited during the academic year 2025-2026.")
doc.add_paragraph("The student has successfully completed the assigned tasks and their performance was found to be satisfactory.")
doc.add_paragraph('\n\n\n')
doc.add_paragraph("Industry Mentor:\n[Mentor Name]\nGrownited")
doc.add_paragraph('\n\n\n')
doc.add_paragraph("Date: ____________________\t\t\t\tSignature: ____________________")
doc.add_page_break()

# 4. Faculty Mentor Certificate
doc.add_heading('Faculty Mentor Certificate', level=1)
doc.add_paragraph("This is to certify that the project entitled \"Stock price prediction system\" submitted by meet siddhapara (AU2240243) is a bonafide work carried out under my supervision and guidance in partial fulfillment of the requirements for the degree of Bachelor of Technology in CSE from the School of Engineering and Applied Science (SEAS), Ahmedabad University.")
doc.add_paragraph('\n\n\n')
doc.add_paragraph("Faculty Mentor:\n[Faculty Name]\nSchool of Engineering and Applied Science (SEAS)\nAhmedabad University")
doc.add_paragraph('\n\n\n')
doc.add_paragraph("Date: ____________________\t\t\t\tSignature: ____________________")
doc.add_page_break()

# 5. Abstract
doc.add_heading('Abstract', level=1)
doc.add_paragraph("This project report details the development of a comprehensive Stock Price Prediction System and Screener application, tailored for the National Stock Exchange (NSE). The system leverages advanced Machine Learning techniques, specifically Long Short-Term Memory (LSTM) neural networks, to forecast next-day stock prices based on extensive historical data. In addition to deep learning predictions, the platform serves as an all-in-one financial assistant for retail investors. It features robust capabilities including real-time fundamental information retrieval, interactive technical indicators (such as MACD, RSI, and Bollinger Bands), an automated stock screener for breakout signals, and candlestick pattern recognition. The project demonstrates the successful integration of a machine learning backend with an intuitive, interactive frontend built using Streamlit. Key outcomes include the creation of a responsive, user-friendly dashboard, an AI-powered trading coach for strategy simulation, and a scalable SQLite database for efficient data management. The system empowers investors with data-driven insights and sophisticated analytical tools to make informed trading decisions.")
doc.add_page_break()

# 6. Table of Contents
doc.add_heading('Table of Contents', level=1)
p = doc.add_paragraph()
run = p.add_run()
fldChar = OxmlElement('w:fldChar')
fldChar.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')

r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar3)
doc.add_paragraph("Note: To update the Table of Contents automatically in Word, right-click on it and select 'Update Field'.")
doc.add_page_break()

# Chapter 1
doc.add_heading('Chapter 1: Introduction', level=1)
doc.add_heading('1.1 | Project Definition', level=2)
doc.add_paragraph("The objective of this project is to build an interactive, data-driven web application that serves as an all-in-one financial assistant for retail investors, specifically focused on National Stock Exchange (NSE) stocks. It integrates fundamental data, technical indicators, pattern recognition, and a deep learning (LSTM) model for next-day price forecasting.")

doc.add_heading('1.2 | Project Objectives', level=2)
p_obj = doc.add_paragraph(style='List Bullet')
p_obj.add_run("Develop an automated stock screener to identify breakout and consolidation zones.")
p_obj = doc.add_paragraph(style='List Bullet')
p_obj.add_run("Implement and fine-tune a Long Short-Term Memory (LSTM) network capable of accurately predicting short-term stock price movements.")
p_obj = doc.add_paragraph(style='List Bullet')
p_obj.add_run("Build a user-friendly frontend dashboard (using Streamlit) with a modern, responsive UI.")
p_obj = doc.add_paragraph(style='List Bullet')
p_obj.add_run("Provide a simulated investment environment (AI Trading Coach) to help users practice trading strategies and receive real-time feedback.")

doc.add_paragraph("\n[Figure 1.1: High-Level System Architecture Diagram]")
doc.add_paragraph("Figure 1.1 illustrates the flow of data from Yahoo Finance API, through the data processing and machine learning layers, and finally to the Streamlit UI.")

doc.add_page_break()

# Chapter 2
doc.add_heading('Chapter 2: Literature Survey', level=1)
doc.add_heading('2.1 | Related Work', level=2)
doc.add_paragraph("Numerous studies have been conducted on stock market prediction utilizing various statistical and machine learning methodologies. Traditional approaches relied heavily on ARIMA and linear regression. More recent advancements leverage deep learning, with Recurrent Neural Networks (RNNs) and their variant, LSTMs, showing significant superiority in capturing the non-linear, temporal dependencies inherent in financial time-series data.")

doc.add_heading('2.2 | Tools and Technologies', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Category'
hdr_cells[1].text = 'Technology Used'
row_cells = table.add_row().cells
row_cells[0].text = 'Programming Language'
row_cells[1].text = 'Python'
row_cells = table.add_row().cells
row_cells[0].text = 'Data Manipulation'
row_cells[1].text = 'Pandas, NumPy, yfinance'
row_cells = table.add_row().cells
row_cells[0].text = 'Machine Learning'
row_cells[1].text = 'TensorFlow, Keras (LSTM), Scikit-Learn'
row_cells = table.add_row().cells
row_cells[0].text = 'Frontend Framework'
row_cells[1].text = 'Streamlit'
row_cells = table.add_row().cells
row_cells[0].text = 'Database'
row_cells[1].text = 'SQLite'
doc.add_paragraph("Table 2.1: Tools and Technologies used in the project.")

doc.add_paragraph("\n[Figure 2.1: Comparison of Traditional vs Deep Learning Approaches]")
doc.add_page_break()

# Chapter 3
doc.add_heading('Chapter 3: Methodology', level=1)
doc.add_paragraph("The methodology for building the Stock Price Prediction System consisted of several discrete phases:")

doc.add_heading('3.1 | Data Collection and Preprocessing', level=2)
doc.add_paragraph("Historical stock data for NSE listed companies was retrieved dynamically using the Yahoo Finance API. Data points included Open, High, Low, Close (OHLC), and Volume metrics covering a 10-year period. Preprocessing involved handling missing values, calculating technical indicators (like Moving Averages), and scaling the features using Min-Max scaling to ensure optimal neural network convergence.")

doc.add_heading('3.2 | Model Architecture', level=2)
doc.add_paragraph("An LSTM network was selected for the predictive core due to its effectiveness in sequence prediction. The architecture involves multiple LSTM layers with Dropout layers interspersed to mitigate overfitting. The model was trained on rolling 60-day windows to predict the subsequent day's closing price.")
doc.add_paragraph("\n[Figure 3.1: LSTM Model Layer Architecture]")

doc.add_heading('3.3 | Project Timeline (Gantt Chart representation)', level=2)
gantt_table = doc.add_table(rows=1, cols=4)
gantt_table.style = 'Table Grid'
hdr_cells = gantt_table.rows[0].cells
hdr_cells[0].text = 'Phase'
hdr_cells[1].text = 'Task Description'
hdr_cells[2].text = 'Duration'
hdr_cells[3].text = 'Completion %'
tasks = [
    ('Phase 1', 'Requirements & Setup', 'Week 1-2', '100%'),
    ('Phase 2', 'Data Retrieval & DB', 'Week 3-5', '100%'),
    ('Phase 3', 'Technical Analysis Module', 'Week 6-8', '100%'),
    ('Phase 4', 'Screener & Patterns', 'Week 9-11', '100%'),
    ('Phase 5', 'LSTM Model Training', 'Week 12-15', '100%'),
    ('Phase 6', 'UI Polish & Testing', 'Week 16-17', '100%')
]
for p, t, d, c in tasks:
    row_cells = gantt_table.add_row().cells
    row_cells[0].text = p
    row_cells[1].text = t
    row_cells[2].text = d
    row_cells[3].text = c
doc.add_paragraph("Table 3.1: Project Timeline and Milestones.")
doc.add_page_break()

# Chapter 4
doc.add_heading('Chapter 4: Results', level=1)
doc.add_heading('4.1 | Project Outcomes', level=2)
doc.add_paragraph("The complete web application successfully loads and evaluates over 1,700 NSE stocks. The LSTM forecasting model demonstrates a high degree of accuracy in capturing the overall trend of stock movements, though volatility spikes remain challenging. The screener efficiently filters stocks based on predefined breakout parameters within seconds.")
doc.add_paragraph("\n[Figure 4.1: Forecasting Module Performance Graph]")

doc.add_heading('4.2 | My Contributions to the Project', level=2)
doc.add_paragraph("During the internship at Grownited, my primary contributions included developing the frontend dashboard, writing the data retrieval pipelines, integrating the technical analysis libraries, and designing the LSTM prediction module from scratch.")

doc.add_heading('4.3 | Learning Outcomes', level=2)
p_lrn = doc.add_paragraph(style='List Bullet')
p_lrn.add_run("Gained deep insights into financial markets and algorithmic trading indicators.")
p_lrn = doc.add_paragraph(style='List Bullet')
p_lrn.add_run("Acquired hands-on experience in building, tuning, and deploying deep learning models using TensorFlow/Keras.")
p_lrn = doc.add_paragraph(style='List Bullet')
p_lrn.add_run("Learned how to develop and structure full-stack data applications using Streamlit.")

doc.add_heading('4.4 | Real World Applications', level=2)
doc.add_paragraph("This application can be utilized by retail investors to conduct rapid, comprehensive fundamental and technical analysis, mitigating the need for expensive premium charting software. The predictive insights also serve as a secondary confirmation tool for trade execution.")
doc.add_page_break()

# Chapter 5
doc.add_heading('Chapter 5: Conclusion', level=1)
doc.add_paragraph("The Stock Price Prediction System bridges the gap between complex machine learning techniques and everyday retail investing. By providing an intuitive interface coupled with deep learning-powered insights, the platform successfully serves its purpose as a comprehensive financial assistant. Future iterations may include real-time intraday data processing, integration of sentiment analysis from financial news feeds, and expanding the model's forecasting horizon to predict multi-day trajectories with confidence intervals.")
doc.add_page_break()

# Chapter 6
doc.add_heading('Chapter 6: Example Chapter', level=1)
doc.add_heading('6.1 | Only for inspiration', level=2)
doc.add_heading('6.1.1 | This is a subsectio', level=3)
doc.add_paragraph("Content for subsection 6.1.1 goes here.")
doc.add_heading('6.2 | General formatting', level=2)
doc.add_paragraph("Content for general formatting goes here.")
doc.add_heading('6.3 | Tables and figures', level=2)
doc.add_paragraph("Content for tables and figures goes here.")
doc.add_page_break()

# Bibliography
doc.add_heading('Bibliography', level=1)
doc.add_paragraph("[1] Hochreiter, S., & Schmidhuber, J. (1997). Long short-term memory. Neural computation, 9(8), 1735-1780.")
doc.add_paragraph("[2] Python Software Foundation. (2026). Python Language Reference. Available at http://www.python.org")
doc.add_paragraph("[3] Abadi, M. et al. (2015). TensorFlow: Large-Scale Machine Learning on Heterogeneous Systems.")
doc.add_page_break()

# Appendix
doc.add_heading('Appendix', level=1)
doc.add_heading('Appendix A title', level=2)
doc.add_paragraph("model = Sequential()\nmodel.add(LSTM(units=50, return_sequences=True, input_shape=(x_train.shape[1], 1)))\nmodel.add(Dropout(0.2))\nmodel.add(LSTM(units=50, return_sequences=False))\nmodel.add(Dropout(0.2))\nmodel.add(Dense(units=25))\nmodel.add(Dense(units=1))\nmodel.compile(optimizer='adam', loss='mean_squared_error')")

doc.save("Stock_Price_Prediction_System_Report.docx")
